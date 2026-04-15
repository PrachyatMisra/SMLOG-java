#!/usr/bin/env python3
from __future__ import annotations

import html
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from PIL import Image as PILImage
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    XPreformatted,
)


ROOT = Path(__file__).resolve().parent
OUTPUT_PDF = ROOT / "SMART_LOGISTICS_MANAGEMENT_SYSTEM_FINAL_SUBMISSION.pdf"
OUTPUT_DOCX = ROOT / "SMART_LOGISTICS_MANAGEMENT_SYSTEM_FINAL_SUBMISSION.docx"
SVG_ARCHITECTURE = ROOT / "report-assets" / "mvc-architecture.svg"
PNG_ARCHITECTURE = ROOT / "report-assets" / "mvc-architecture-report.png"


TEAM_MEMBERS = [
    {
        "name": "Sujal Purbey",
        "roll": "25MCA1035",
        "responsibility": "Frontend implementation, browser dashboard, and JavaFX desktop interface.",
    },
    {
        "name": "Sudhanshu Ranjan",
        "roll": "25MCA1053",
        "responsibility": "Machine learning logic, Python analytics module, and ML integration pipeline.",
    },
    {
        "name": "Prachyat Misra",
        "roll": "25MCA1063",
        "responsibility": "Overall integration, backend architecture, servlet APIs, DAO layer, and submission assembly.",
    },
]


CORE_FILES = [
    "backend/pom.xml",
    "backend/src/main/java/com/smartlogistics/SmartLogisticsApplication.java",
    "backend/src/main/java/com/smartlogistics/dao/DBConnection.java",
    "backend/src/main/java/com/smartlogistics/dao/ShipmentDAO.java",
    "backend/src/main/java/com/smartlogistics/ml/AnomalyDetector.java",
    "backend/src/main/java/com/smartlogistics/ml/DemandPredictor.java",
    "backend/src/main/java/com/smartlogistics/ml/MLService.java",
    "backend/src/main/java/com/smartlogistics/model/Shipment.java",
    "backend/src/main/java/com/smartlogistics/servlet/AnomalyServlet.java",
    "backend/src/main/java/com/smartlogistics/servlet/BaseApiServlet.java",
    "backend/src/main/java/com/smartlogistics/servlet/ForecastServlet.java",
    "backend/src/main/java/com/smartlogistics/servlet/ShipmentServlet.java",
    "backend/src/main/resources/application.properties",
    "backend/src/main/resources/database.sql",
    "backend/src/main/webapp/WEB-INF/web.xml",
    "javafx-frontend/pom.xml",
    "javafx-frontend/src/main/java/com/smartlogistics/javafx/SmartLogisticsApp.java",
    "ml/ml_predictor.py",
    "mysql-workbench-demo.sql",
    "start-project.sh",
    "stop-project.sh",
    "web/index.html",
    "workbench-all-shipments.sql",
    "workbench-demo.sql",
    "workbench_refresh_grid.py",
]


SCREENSHOTS = [
    ("Browser dashboard overview", ROOT / "report-assets" / "dashboard-running.png", "Live servlet-backed dashboard interface."),
    ("Dashboard CRUD operations", ROOT / "report-assets" / "dashboard-crud-before-delete.png", "Shipment management form and summary cards."),
    ("Shipment API list endpoint", ROOT / "report-assets" / "chrome-shipments-live.png", "Servlet response for all shipment records."),
    ("Forecast endpoint output", ROOT / "report-assets" / "chrome-forecast-live.png", "Demand prediction returned through the ML servlet."),
    ("Anomaly endpoint output", ROOT / "report-assets" / "chrome-anomalies-live.png", "Delayed shipment anomalies surfaced by the analytics module."),
    ("Database evidence", ROOT / "report-assets" / "workbench-live-text.png", "MySQL Workbench showing persistent shipment records."),
    ("JavaFX desktop client", ROOT / "report-assets" / "javafx-running.png", "Desktop interface consuming the same backend APIs."),
    ("Dashboard after delete", ROOT / "report-assets" / "dashboard-after-delete.png", "Frontend state after servlet-backed deletion flow."),
]


TEAM_PORTFOLIOS = [
    {
        "name": "Sujal Purbey",
        "roll": "25MCA1035",
        "role": "Frontend Engineering, Browser Dashboard, and JavaFX Desktop Experience",
        "summary": "Sujal Purbey was responsible for the user-facing layers of the project. His work delivered the responsive browser dashboard and the JavaFX desktop application that consume the same live shipment APIs, making the project accessible in both web and desktop form factors.",
        "files": [
            "web/index.html",
            "javafx-frontend/src/main/java/com/smartlogistics/javafx/SmartLogisticsApp.java",
            "javafx-frontend/pom.xml",
        ],
        "screenshots": [
            ("Dashboard interface", ROOT / "report-assets" / "dashboard-running.png", "Browser dashboard showing the main project interface."),
            ("Dashboard CRUD workflow", ROOT / "report-assets" / "dashboard-crud-before-delete.png", "Shipment creation and monitoring flow on the web interface."),
            ("JavaFX desktop client", ROOT / "report-assets" / "javafx-running.png", "Desktop interface built for the same project backend."),
        ],
        "data_rows": [
            ["Primary UI channels", "Web dashboard and JavaFX desktop client"],
            ["User-facing capabilities", "Create, view, edit, delete, stats refresh, forecast trigger, anomaly trigger"],
            ["Main integration style", "Frontend requests to shared servlet endpoints over HTTP"],
            ["Submission evidence", "Dashboard screenshots, JavaFX screenshot, UI code excerpts"],
        ],
    },
    {
        "name": "Sudhanshu Ranjan",
        "roll": "25MCA1053",
        "role": "Machine Learning Logic and Java-Python Analytics Integration",
        "summary": "Sudhanshu Ranjan implemented the AI-assisted layer of the project. His contribution includes the Python analytics script, demand prediction logic, anomaly detection logic, and the backend bridge that converts servlet requests into analytics outputs returned as JSON.",
        "files": [
            "ml/ml_predictor.py",
            "backend/src/main/java/com/smartlogistics/ml/MLService.java",
            "backend/src/main/java/com/smartlogistics/ml/DemandPredictor.java",
            "backend/src/main/java/com/smartlogistics/ml/AnomalyDetector.java",
            "backend/src/main/java/com/smartlogistics/servlet/ForecastServlet.java",
            "backend/src/main/java/com/smartlogistics/servlet/AnomalyServlet.java",
        ],
        "screenshots": [
            ("Forecast endpoint", ROOT / "report-assets" / "chrome-forecast-live.png", "Prediction response returned through the analytics endpoint."),
            ("Anomaly endpoint", ROOT / "report-assets" / "chrome-anomalies-live.png", "Delayed-shipment anomalies returned through the analytics endpoint."),
            ("Servlet forecast view", ROOT / "report-assets" / "servlet-forecast-endpoint.png", "Servlet presentation of machine learning output."),
        ],
        "data_rows": [
            ["Analytics operations", "Demand prediction and delayed-shipment anomaly detection"],
            ["Input data source", "Shipment count and shipment list from backend DAO results"],
            ["Integration method", "Java ProcessBuilder invoking Python with JSON exchange"],
            ["Output format", "Structured JSON returned to servlet and frontend layers"],
        ],
    },
    {
        "name": "Prachyat Misra",
        "roll": "25MCA1063",
        "role": "Backend Architecture, Database Layer, and Overall Project Integration",
        "summary": "Prachyat Misra led the system integration and backend implementation. This work includes servlet architecture, application bootstrap, shipment CRUD routing, JDBC persistence, database schema setup, startup execution flow, and final consolidation of the project deliverables for submission.",
        "files": [
            "backend/src/main/java/com/smartlogistics/SmartLogisticsApplication.java",
            "backend/src/main/java/com/smartlogistics/dao/DBConnection.java",
            "backend/src/main/java/com/smartlogistics/dao/ShipmentDAO.java",
            "backend/src/main/java/com/smartlogistics/model/Shipment.java",
            "backend/src/main/java/com/smartlogistics/servlet/BaseApiServlet.java",
            "backend/src/main/java/com/smartlogistics/servlet/ShipmentServlet.java",
            "backend/src/main/resources/database.sql",
            "backend/src/main/webapp/WEB-INF/web.xml",
            "start-project.sh",
            "stop-project.sh",
        ],
        "screenshots": [
            ("Shipment servlet response", ROOT / "report-assets" / "chrome-shipments-live.png", "Backend shipment list endpoint working live."),
            ("Database persistence evidence", ROOT / "report-assets" / "workbench-live-text.png", "MySQL Workbench showing persisted shipment rows."),
            ("Single shipment retrieval", ROOT / "report-assets" / "servlet-single-shipment.png", "Single-record backend retrieval through servlet routing."),
        ],
        "data_rows": [
            ["Core responsibility", "System-wide integration, backend orchestration, persistence, and deployment preparation"],
            ["Backend endpoints", "Shipment CRUD plus analytics delegation endpoints"],
            ["Database scope", "Schema creation, seed data, CRUD access, and live persistence"],
            ["Operational support", "Project startup, stop scripts, and final packaging"],
        ],
    },
]


@dataclass(frozen=True)
class SnippetSpec:
    title: str
    path: str
    start_marker: str
    end_marker: str | None
    max_lines: int
    commentary: str


SNIPPETS = [
    SnippetSpec(
        "Backend schema bootstrap",
        "backend/src/main/java/com/smartlogistics/dao/DBConnection.java",
        "private static void ensureSchema()",
        "private static boolean isSeedDataPresent()",
        42,
        "The backend initializes the MySQL database, creates the shipments table if required, and seeds starter records for live demonstration.",
    ),
    SnippetSpec(
        "Shipment CRUD servlet",
        "backend/src/main/java/com/smartlogistics/servlet/ShipmentServlet.java",
        "protected void doGet",
        "private String extractShipmentId",
        80,
        "The servlet layer exposes create, read, update, and delete operations and returns JSON to both frontends.",
    ),
    SnippetSpec(
        "Java to Python ML bridge",
        "backend/src/main/java/com/smartlogistics/ml/MLService.java",
        "public static Map<String, Object> predictDemand",
        "private static Path resolveScriptPath()",
        34,
        "Machine learning is integrated through a Java service that invokes the Python analytics script and forwards JSON results.",
    ),
    SnippetSpec(
        "Python analytics logic",
        "ml/ml_predictor.py",
        "def predict_demand",
        "def main():",
        45,
        "The Python module performs demand prediction and delayed-shipment anomaly detection.",
    ),
    SnippetSpec(
        "Web dashboard rendering",
        "web/index.html",
        "async function loadShipments()",
        "async function loadShipmentById()",
        52,
        "The browser dashboard dynamically fetches shipment data, updates the statistics cards, and renders shipment cards.",
    ),
    SnippetSpec(
        "JavaFX desktop integration",
        "javafx-frontend/src/main/java/com/smartlogistics/javafx/SmartLogisticsApp.java",
        "private void loadShipments()",
        "private HBox createStatsPanel",
        80,
        "The JavaFX client consumes the same servlet API and presents shipment operations in a desktop environment.",
    ),
]


def sanitize_text(text: str) -> str:
    text = text.replace("Harshbabu@2004", "[REDACTED_PASSWORD]")
    text = re.sub(
        r'(SMART_LOGISTICS_DB_PASSWORD",\s*")[^"]+(")',
        r"\1[REDACTED_PASSWORD]\2",
        text,
    )
    return text


def get_repo_url() -> str | None:
    try:
        raw = subprocess.check_output(
            ["git", "remote", "get-url", "origin"],
            cwd=ROOT,
            text=True,
            stderr=subprocess.DEVNULL,
        ).strip()
    except Exception:
        return None

    if raw.startswith("git@github.com:"):
        raw = raw.replace("git@github.com:", "https://github.com/")
    raw = raw.removesuffix(".git")
    if raw.startswith("https://github.com/"):
        return raw
    return raw if raw.startswith("http") else None


def owner_for_path(rel_path: str) -> str:
    if rel_path.startswith("web/") or rel_path.startswith("javafx-frontend/"):
        return "Sujal Purbey (Frontend and JavaFX)"
    if rel_path.startswith("ml/") or "/ml/" in rel_path or rel_path.endswith("ForecastServlet.java") or rel_path.endswith("AnomalyServlet.java"):
        return "Sudhanshu Ranjan (ML and integration)"
    return "Prachyat Misra (Backend and overall integration)"


def language_for_path(path: str) -> str:
    suffix = Path(path).suffix.lower()
    return {
        ".java": "Java",
        ".py": "Python",
        ".html": "HTML/CSS/JS",
        ".xml": "XML",
        ".sql": "SQL",
        ".sh": "Shell",
        ".properties": "Properties",
    }.get(suffix, "Text")


def load_code(rel_path: str) -> str:
    return sanitize_text((ROOT / rel_path).read_text(encoding="utf-8", errors="ignore"))


def count_loc(rel_path: str) -> int:
    return len(load_code(rel_path).splitlines())


def count_loc_many(paths: list[str]) -> int:
    return sum(count_loc(path) for path in paths)


def extract_snippet(spec: SnippetSpec) -> str:
    lines = load_code(spec.path).splitlines()
    start_index = 0
    for idx, line in enumerate(lines):
        if spec.start_marker in line:
            start_index = idx
            break
    end_index = min(len(lines), start_index + spec.max_lines)
    if spec.end_marker:
        for idx in range(start_index + 1, len(lines)):
            if spec.end_marker in lines[idx]:
                end_index = min(idx, start_index + spec.max_lines)
                break
    return "\n".join(lines[start_index:end_index]).rstrip()


def code_block(text: str, style: ParagraphStyle) -> XPreformatted:
    return XPreformatted(html.escape(text.rstrip()), style)


def ensure_architecture_png() -> Path:
    if PNG_ARCHITECTURE.exists():
        return PNG_ARCHITECTURE
    try:
        subprocess.run(
            ["magick", str(SVG_ARCHITECTURE), str(PNG_ARCHITECTURE)],
            check=True,
            cwd=ROOT,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        return PNG_ARCHITECTURE
    except subprocess.CalledProcessError:
        pass

    preview_dir = ROOT / "report-assets"
    try:
        subprocess.run(
            [
                "qlmanage",
                "-t",
                "-s",
                "1800",
                "-o",
                str(preview_dir),
                str(SVG_ARCHITECTURE),
            ],
            check=True,
            cwd=ROOT,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        generated = preview_dir / f"{SVG_ARCHITECTURE.name}.png"
        if generated.exists():
            generated.replace(PNG_ARCHITECTURE)
            return PNG_ARCHITECTURE
    except subprocess.CalledProcessError:
        pass

    sample_page = ROOT / "report-assets" / "sample-pages" / "sample-page-3.png"
    if sample_page.exists():
        return sample_page

    raise FileNotFoundError("Unable to prepare an architecture image for the report.")


def image_flowable(path: Path, max_width_cm: float = 16.8, max_height_cm: float = 18.0) -> Image:
    img = PILImage.open(path)
    width, height = img.size
    scale = min((max_width_cm * cm) / width, (max_height_cm * cm) / height, 1.0)
    return Image(str(path), width=width * scale, height=height * scale)


def link_markup(url: str | None) -> str:
    if not url:
        return "Repository link will appear here after GitHub publication."
    safe = html.escape(url)
    return f"<link href='{safe}' color='blue'>{safe}</link>"


def build_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="SectionTitle",
            parent=styles["Heading1"],
            fontName="Times-Bold",
            fontSize=20,
            leading=24,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#18324f"),
            spaceBefore=8,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SubHeading",
            parent=styles["Heading2"],
            fontName="Times-Bold",
            fontSize=15,
            leading=18,
            textColor=colors.HexColor("#214e7a"),
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            textColor=colors.HexColor("#1f2f41"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Caption",
            parent=styles["BodyText"],
            fontName="Times-Italic",
            fontSize=9,
            leading=12,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#5b6b7f"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CodeFileTitle",
            parent=styles["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1a3857"),
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CodeBlock",
            fontName="Courier",
            fontSize=6.4,
            leading=7.4,
            leftIndent=6,
            rightIndent=6,
            backColor=colors.HexColor("#f4f7fb"),
            borderColor=colors.HexColor("#d5dee9"),
            borderWidth=0.6,
            borderPadding=6,
            borderRadius=None,
            spaceAfter=8,
            textColor=colors.HexColor("#203246"),
        )
    )
    return styles


def cover_table() -> Table:
    data = [
        ["Report Type", "Final Project Review / DA2 Submission"],
        ["Prepared For", "Faculty review submission based on the original Smart Logistics template"],
        ["Implementation Stage", "Project completed and consolidated for final submission"],
        ["Date", "15 April 2026"],
    ]
    table = Table(data, colWidths=[5.0 * cm, 10.7 * cm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f7fbff")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#1d2f45")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("LEADING", (0, 0), (-1, -1), 14),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#ccd7e5")),
                ("INNERGRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#d6deea")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return table


def simple_table(rows: list[list[str]], col_widths: list[float], font_size: float = 9.4, header_bg: str = "#edf3fb") -> Table:
    table = Table(rows, colWidths=[width * cm for width in col_widths], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(header_bg)),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), font_size),
                ("LEADING", (0, 0), (-1, -1), max(11, font_size + 3)),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#cad5e2")),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d4deea")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


def portfolio_snippet_specs(name: str) -> list[SnippetSpec]:
    mapping = {
        "Sujal Purbey": [
            SnippetSpec(
                "Web Dashboard Shipment Rendering",
                "web/index.html",
                "async function loadShipments()",
                "async function saveShipment()",
                46,
                "The browser interface dynamically fetches shipment data and renders an operational dashboard.",
            ),
            SnippetSpec(
                "JavaFX Shipment Loading Flow",
                "javafx-frontend/src/main/java/com/smartlogistics/javafx/SmartLogisticsApp.java",
                "private void loadShipments()",
                "private HBox createStatsPanel",
                78,
                "The desktop client presents the same backend data through a JavaFX-based shipment interface.",
            ),
        ],
        "Sudhanshu Ranjan": [
            SnippetSpec(
                "Python Demand and Anomaly Logic",
                "ml/ml_predictor.py",
                "def predict_demand",
                "def main():",
                45,
                "The Python module implements both analytics functions used by the application.",
            ),
            SnippetSpec(
                "Java-Python Integration Bridge",
                "backend/src/main/java/com/smartlogistics/ml/MLService.java",
                "public static Map<String, Object> predictDemand",
                "private static Path resolveScriptPath()",
                34,
                "The Java analytics bridge launches the Python script and returns JSON output back to the servlets.",
            ),
        ],
        "Prachyat Misra": [
            SnippetSpec(
                "Database Schema and Seed Bootstrap",
                "backend/src/main/java/com/smartlogistics/dao/DBConnection.java",
                "private static void ensureSchema()",
                "private static boolean isSeedDataPresent()",
                42,
                "The backend ensures schema availability and initializes persistent sample data for the logistics system.",
            ),
            SnippetSpec(
                "Shipment CRUD Routing",
                "backend/src/main/java/com/smartlogistics/servlet/ShipmentServlet.java",
                "protected void doGet",
                "private String extractShipmentId",
                84,
                "The servlet handles CRUD routing, validation, and JSON responses for shipment operations.",
            ),
        ],
    }
    return mapping[name]


def add_team_portfolio(story: list, styles, portfolio: dict, repo_url: str | None) -> None:
    owned_files = portfolio["files"]
    story.append(Paragraph(portfolio["name"], styles["SectionTitle"]))
    story.append(Paragraph(portfolio["role"], styles["SubHeading"]))
    story.append(Paragraph(portfolio["summary"], styles["Body"]))

    summary_rows = [
        ["Field", "Details"],
        ["Roll Number", portfolio["roll"]],
        ["Primary ownership", portfolio["role"]],
        ["Files owned in submission", str(len(owned_files))],
        ["Approximate code volume", f"{count_loc_many(owned_files)} lines across owned modules"],
        ["Public repository access", repo_url or "To be inserted after GitHub publication"],
    ]
    story.append(simple_table(summary_rows, [4.2, 11.0], font_size=9.2))
    story.append(Spacer(1, 0.12 * cm))

    file_rows = [["Key file", "Language", "Lines", "Why it matters"]]
    rationale_map = {
        "web/index.html": "Main browser dashboard and UI behavior.",
        "javafx-frontend/src/main/java/com/smartlogistics/javafx/SmartLogisticsApp.java": "Desktop client for shipment operations.",
        "javafx-frontend/pom.xml": "JavaFX build and execution setup.",
        "ml/ml_predictor.py": "Analytics script for demand prediction and anomalies.",
        "backend/src/main/java/com/smartlogistics/ml/MLService.java": "Backend bridge connecting Java to Python analytics.",
        "backend/src/main/java/com/smartlogistics/ml/DemandPredictor.java": "Forecasting helper logic for demand modelling.",
        "backend/src/main/java/com/smartlogistics/ml/AnomalyDetector.java": "Detection logic for outlier shipments.",
        "backend/src/main/java/com/smartlogistics/servlet/ForecastServlet.java": "Forecast endpoint exposed to clients.",
        "backend/src/main/java/com/smartlogistics/servlet/AnomalyServlet.java": "Anomaly endpoint exposed to clients.",
        "backend/src/main/java/com/smartlogistics/SmartLogisticsApplication.java": "Application bootstrap for the backend.",
        "backend/src/main/java/com/smartlogistics/dao/DBConnection.java": "Database initialization and connection management.",
        "backend/src/main/java/com/smartlogistics/dao/ShipmentDAO.java": "CRUD operations over persistent shipment data.",
        "backend/src/main/java/com/smartlogistics/model/Shipment.java": "Core data model shared across layers.",
        "backend/src/main/java/com/smartlogistics/servlet/BaseApiServlet.java": "Shared JSON and request handling utilities.",
        "backend/src/main/java/com/smartlogistics/servlet/ShipmentServlet.java": "Main shipment CRUD servlet.",
        "backend/src/main/resources/database.sql": "DDL and sample shipment dataset.",
        "backend/src/main/webapp/WEB-INF/web.xml": "Servlet mapping and deployment metadata.",
        "start-project.sh": "Project startup orchestration.",
        "stop-project.sh": "Runtime shutdown support.",
    }
    for path in owned_files:
        file_rows.append([path, language_for_path(path), str(count_loc(path)), rationale_map.get(path, "Project contribution file.")])
    story.append(Paragraph("Owned Source and Deliverables", styles["SubHeading"]))
    story.append(simple_table(file_rows, [7.0, 2.0, 1.3, 5.0], font_size=7.8))

    data_rows = [["Evidence Type", "Details"]] + portfolio["data_rows"]
    story.append(Paragraph("Contribution Data and Deliverable Evidence", styles["SubHeading"]))
    story.append(simple_table(data_rows, [4.6, 10.6], font_size=8.8))

    story.append(Paragraph("Dedicated Screenshots", styles["SubHeading"]))
    for title, path, caption in portfolio["screenshots"]:
        story.append(Paragraph(title, styles["CodeFileTitle"]))
        story.append(image_flowable(path, max_width_cm=15.8, max_height_cm=10.4))
        story.append(Paragraph(caption, styles["Caption"]))

    story.append(Paragraph("Dedicated Code Evidence", styles["SubHeading"]))
    for spec in portfolio_snippet_specs(portfolio["name"]):
        story.append(Paragraph(spec.title, styles["CodeFileTitle"]))
        story.append(Paragraph(spec.commentary, styles["Body"]))
        story.append(code_block(extract_snippet(spec), styles["CodeBlock"]))

    story.append(PageBreak())


def build_story() -> list:
    styles = build_styles()
    architecture_png = ensure_architecture_png()
    total_loc = sum(count_loc(path) for path in CORE_FILES)
    repo_url = get_repo_url()

    story = []

    story.append(Spacer(1, 4.0 * cm))
    story.append(
        Paragraph(
            "PROGRESS REVIEW 3 / FINAL SUBMISSION",
            ParagraphStyle(
                "CoverTop",
                fontName="Helvetica-Bold",
                fontSize=13,
                leading=16,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#214e7a"),
                tracking=1.5,
                spaceAfter=16,
            ),
        )
    )
    story.append(
        Paragraph(
            "SMART LOGISTICS MANAGEMENT SYSTEM",
            ParagraphStyle(
                "CoverTitle",
                fontName="Times-Bold",
                fontSize=24,
                leading=31,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#18324f"),
                spaceAfter=10,
            ),
        )
    )
    story.append(
        Paragraph(
            "(SMLOG-JAVA)",
            ParagraphStyle(
                "CoverCode",
                fontName="Helvetica-Bold",
                fontSize=16,
                leading=18,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#18324f"),
                spaceAfter=12,
            ),
        )
    )
    story.append(
        Paragraph(
            "Servlet-based shipment tracking platform with MySQL persistence, browser dashboard, JavaFX desktop client, and Python machine learning integration.",
            ParagraphStyle(
                "CoverSubtitle",
                fontName="Times-Roman",
                fontSize=12,
                leading=18,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#4d6178"),
                spaceAfter=18,
            ),
        )
    )
    story.append(cover_table())
    story.append(Spacer(1, 0.5 * cm))
    story.append(
        Paragraph(
            f"<b>Public GitHub Repository:</b> {link_markup(repo_url)}",
            styles["Body"],
        )
    )
    story.append(
        Paragraph(
            "This submission professionally consolidates the final project state using the earlier review style, updated Review 3 status, richer content structure, execution evidence, detailed team contribution portfolios, public repository access, and a full source-code appendix.",
            styles["Body"],
        )
    )
    story.append(PageBreak())

    story.append(Paragraph("Project Overview", styles["SectionTitle"]))
    story.append(
        Paragraph(
            "Smart Logistics Management System is a multi-layer Java project designed to manage shipments, expose servlet-backed APIs, persist logistics data in MySQL, offer both browser and JavaFX interfaces, and provide AI-assisted forecasting and anomaly detection through Python integration.",
            styles["Body"],
        )
    )
    story.append(
        Paragraph(
            "The current submission is a continuation of the earlier Progress Review 1 and 2 document. In this final version, the project has been organized as a complete working system with demonstrable CRUD operations, analytics endpoints, integrated frontends, and a documented appendix containing the real project code submitted for evaluation.",
            styles["Body"],
        )
    )
    project_rows = [
        ["Project Attribute", "Details"],
        ["Project title", "Smart Logistics Management System (SMLOG-JAVA)"],
        ["Prepared for", "Final academic project submission / Progress Review 3"],
        ["Public repository", repo_url or "Repository link pending publication"],
        ["Team strength", "3 members"],
        ["Major modules", "Backend, database, browser dashboard, JavaFX client, ML analytics, execution scripts"],
        ["Verified submission outputs", "Final PDF, final DOCX, public repository, source-code appendix"],
    ]
    story.append(Paragraph("Project and Access Information", styles["SubHeading"]))
    story.append(simple_table(project_rows, [4.6, 10.6], font_size=9.0))
    story.append(Paragraph("Objectives", styles["SubHeading"]))
    for item in [
        "Develop a smart shipment management system using Java, servlets, JDBC, and MySQL.",
        "Provide a responsive browser dashboard for shipment creation, editing, monitoring, and analytics access.",
        "Provide a JavaFX desktop application that consumes the same backend services.",
        "Integrate a Python analytics layer for demand prediction and delayed-shipment anomaly detection.",
        "Deliver a professional, evidence-backed submission with screenshots, team contributions, and source-code appendix.",
    ]:
        story.append(Paragraph(f"• {item}", styles["Body"]))
    story.append(Paragraph("Technology Stack", styles["SubHeading"]))
    tech_table = Table(
        [
            ["Layer", "Technology"],
            ["Backend", "Java 17, Spring Boot runtime, Jakarta Servlets, Gson"],
            ["Persistence", "MySQL, JDBC, SQL schema bootstrap and seed data"],
            ["Browser Frontend", "HTML, CSS, vanilla JavaScript"],
            ["Desktop Frontend", "JavaFX 21, JSON parsing, HTTP client"],
            ["Analytics", "Python 3, command-line ML bridge, JSON exchange"],
            ["Build and Execution", "Maven, shell startup scripts, HTTP server"],
        ],
        colWidths=[4.7 * cm, 11.0 * cm],
    )
    tech_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#edf3fb")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.4),
                ("LEADING", (0, 0), (-1, -1), 13),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#cad5e2")),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d4deea")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(tech_table)
    story.append(PageBreak())

    story.append(Paragraph("Team Contributions", styles["SectionTitle"]))
    contribution_rows = [["Team Member", "Roll Number", "Primary Responsibility", "Public Access Reference"]]
    for member in TEAM_MEMBERS:
        contribution_rows.append([member["name"], member["roll"], member["responsibility"], repo_url or "Pending GitHub publication"])
    contribution_table = simple_table(contribution_rows, [3.6, 2.4, 6.0, 3.2], font_size=8.4)
    story.append(contribution_table)
    story.append(Spacer(1, 0.2 * cm))
    for paragraph in [
        "The contribution model for this report has been expanded beyond the original combined table so that each team member now has a dedicated evidence portfolio with owned files, screenshots, data notes, and code excerpts.",
        "This format makes the submission more precise for academic review because it clearly separates frontend, machine learning, and backend/system-integration responsibilities while still presenting the project as one integrated system.",
    ]:
        story.append(Paragraph(paragraph, styles["Body"]))
    story.append(PageBreak())

    story.append(Paragraph("Individual Contribution Portfolios", styles["SectionTitle"]))
    story.append(
        Paragraph(
            "The following sections present individual member-wise contributions with dedicated screenshots, contribution data, and code excerpts mapped directly to the files submitted in the project.",
            styles["Body"],
        )
    )
    story.append(PageBreak())
    for portfolio in TEAM_PORTFOLIOS:
        add_team_portfolio(story, styles, portfolio, repo_url)

    story.append(Paragraph("Architecture and Modules", styles["SectionTitle"]))
    story.append(image_flowable(architecture_png, max_width_cm=16.5, max_height_cm=12.5))
    story.append(Paragraph("Figure 1. MVC architecture with the major files grouped by responsibility.", styles["Caption"]))
    story.append(Paragraph("Architecture Explanation", styles["SubHeading"]))
    for item in [
        "<b>View Layer:</b> <font face='Courier'>web/index.html</font> provides the browser dashboard, while <font face='Courier'>SmartLogisticsApp.java</font> provides the desktop GUI.",
        "<b>Controller Layer:</b> The servlet classes expose shipment CRUD operations and analytics endpoints for forecasting and anomaly detection.",
        "<b>Model and DAO Layer:</b> Shipment entities, JDBC connections, and SQL operations ensure persistent logistics data management.",
        "<b>Analytics and Integration Layer:</b> The Java ML bridge invokes the Python script, receives JSON output, and forwards results to both frontends.",
    ]:
        story.append(Paragraph(item, styles["Body"]))

    module_text = [
        (
            "Module 1. Backend and API Layer",
            "The backend layer exposes REST-like servlet endpoints for listing, creating, updating, and deleting shipment records. It also coordinates error handling, JSON serialization, and shared API utilities.",
        ),
        (
            "Module 2. Database and Persistence Layer",
            "The persistence layer connects the application with MySQL, creates the schema if required, seeds demonstration data, and performs CRUD actions through JDBC.",
        ),
        (
            "Module 3. Browser Dashboard",
            "The dashboard provides a visually organized browser interface for shipment entry, shipment monitoring, statistics, and analytics triggers.",
        ),
        (
            "Module 4. Machine Learning and Analytics",
            "The analytics module predicts demand and detects delayed-shipment anomalies using a Python script invoked from Java through a controlled process bridge.",
        ),
        (
            "Module 5. JavaFX Desktop Frontend",
            "The desktop client offers an alternative interface for the same live backend, proving that the application logic can be reused beyond the web dashboard.",
        ),
    ]
    for title, description in module_text:
        story.append(Paragraph(title, styles["SubHeading"]))
        story.append(Paragraph(description, styles["Body"]))
    story.append(PageBreak())

    story.append(Paragraph("Operational Data and System Evidence", styles["SectionTitle"]))
    story.append(
        Paragraph(
            "The system operates on shipment records stored in MySQL and exposed through servlet endpoints. The following table summarizes sample operational records used to demonstrate project functionality in the report.",
            styles["Body"],
        )
    )
    shipment_rows = [
        ["Shipment ID", "Origin", "Destination", "Cargo", "Weight", "Status"],
        ["SHP-1001", "Mumbai", "Delhi", "Electronics", "420.00", "In Transit"],
        ["SHP-1002", "Chennai", "Bangalore", "Pharmaceuticals", "180.00", "Delivered"],
        ["SHP-1003", "Kolkata", "Hyderabad", "Textiles", "650.00", "Pending"],
        ["SHP-1004", "Pune", "Ahmedabad", "Auto Parts", "900.00", "Delayed"],
        ["SHP-1005", "Jaipur", "Surat", "FMCG Goods", "300.00", "In Transit"],
    ]
    story.append(simple_table(shipment_rows, [2.4, 2.3, 2.7, 3.0, 2.0, 2.6], font_size=8.4))
    analytics_rows = [
        ["Analytics Output", "Interpretation"],
        ["Predicted demand next week", "Machine learning output generated from current shipment count"],
        ["Anomaly count", "Number of delayed shipments requiring attention"],
        ["Anomaly rate percentage", "Delayed shipment ratio relative to total shipment records"],
        ["System health", "Operational signal derived from the anomaly rate"],
    ]
    story.append(Paragraph("Machine Learning Output Fields", styles["SubHeading"]))
    story.append(simple_table(analytics_rows, [4.8, 10.4], font_size=8.8))
    story.append(PageBreak())

    story.append(Paragraph("Execution Evidence", styles["SectionTitle"]))
    for index, (title, path, caption) in enumerate(SCREENSHOTS, start=1):
        story.append(Paragraph(title, styles["SubHeading"]))
        story.append(image_flowable(path, max_width_cm=16.4, max_height_cm=14.2))
        story.append(Paragraph(f"Figure {index + 1}. {caption}", styles["Caption"]))
    story.append(PageBreak())

    story.append(Paragraph("Completion Status", styles["SectionTitle"]))
    status_table = Table(
        [
            ["Module", "Outcome", "Status", "Owner"],
            ["Backend servlet layer", "Shipment CRUD endpoints and shared JSON utilities", "Completed", "Prachyat Misra"],
            ["Database layer", "MySQL schema, seed data, and JDBC DAO flow", "Completed", "Prachyat Misra"],
            ["Browser dashboard", "Operational dashboard with forms, cards, stats, and API integration", "Completed", "Sujal Purbey"],
            ["Machine learning", "Demand prediction and anomaly detection pipeline", "Completed", "Sudhanshu Ranjan"],
            ["Desktop frontend", "JavaFX client connected to live APIs", "Completed", "Sujal Purbey"],
            ["System integration", "End-to-end coordination and final submission packaging", "Completed", "Prachyat Misra"],
        ],
        colWidths=[4.3 * cm, 6.2 * cm, 2.1 * cm, 3.4 * cm],
    )
    status_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#edf3fb")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.9),
                ("LEADING", (0, 0), (-1, -1), 12),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#cad5e2")),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d4deea")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(status_table)
    story.append(Spacer(1, 0.2 * cm))
    story.append(
        Paragraph(
            "The codebase analyzed for this report contains "
            f"<b>{len(CORE_FILES)}</b> submitted source, configuration, and execution-support files with a combined size of "
            f"<b>{total_loc}</b> lines. The appendix below includes the complete textual project contents used for the submission package.",
            styles["Body"],
        )
    )
    story.append(
        Paragraph(
            "Sensitive local credentials found in the repository were masked only inside the report appendix for safe academic submission. The project files themselves were not modified.",
            styles["Body"],
        )
    )
    story.append(
        Paragraph(
            "The public repository link has been placed in both the project information and team information sections so the evaluator can verify source access directly from the report.",
            styles["Body"],
        )
    )
    story.append(Paragraph("Conclusion", styles["SubHeading"]))
    story.append(
        Paragraph(
            "The Smart Logistics Management System has been completed as a multi-interface, multi-layer application that demonstrates logistics data management, real-time interface design, backend persistence, and AI-assisted insight generation. The final submission now reflects stronger academic presentation quality through richer layout, clearer responsibility mapping, separate evidence per team member, and public repository availability.",
            styles["Body"],
        )
    )
    story.append(PageBreak())

    story.append(Paragraph("Appendix A. Source Code Inventory", styles["SectionTitle"]))
    inventory_rows = [["File", "Language", "Lines", "Primary Responsibility"]]
    for rel_path in CORE_FILES:
        inventory_rows.append([rel_path, language_for_path(rel_path), str(count_loc(rel_path)), owner_for_path(rel_path)])
    inventory_table = Table(inventory_rows, colWidths=[7.0 * cm, 2.4 * cm, 1.5 * cm, 5.1 * cm], repeatRows=1)
    inventory_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#edf3fb")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.6),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("BOX", (0, 0), (-1, -1), 0.55, colors.HexColor("#cad5e2")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d4deea")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(inventory_table)
    story.append(PageBreak())

    story.append(Paragraph("Appendix B. Full Source Code Listings", styles["SectionTitle"]))
    story.append(
        Paragraph(
            "The following pages include the full textual source submitted from the SMLOG-JAVA project so that the report remains self-contained for academic review.",
            styles["Body"],
        )
    )
    for rel_path in CORE_FILES:
        code = load_code(rel_path)
        metadata = f"Path: {rel_path} | Language: {language_for_path(rel_path)} | Lines: {len(code.splitlines())} | Owner: {owner_for_path(rel_path)}"
        story.append(Paragraph(rel_path, styles["CodeFileTitle"]))
        story.append(Paragraph(metadata, styles["Caption"]))
        story.append(code_block(code, styles["CodeBlock"]))

    return story


def draw_pdf_header_footer(canvas, doc):
    canvas.saveState()
    width, height = A4
    canvas.setStrokeColor(colors.HexColor("#cfd8e3"))
    canvas.setLineWidth(0.6)
    canvas.line(1.8 * cm, height - 1.5 * cm, width - 1.8 * cm, height - 1.5 * cm)
    canvas.line(1.8 * cm, 1.3 * cm, width - 1.8 * cm, 1.3 * cm)
    canvas.setFont("Helvetica", 8.5)
    canvas.setFillColor(colors.HexColor("#586b80"))
    canvas.drawString(1.8 * cm, height - 1.15 * cm, "Smart Logistics Management System | Final Submission")
    canvas.drawRightString(width - 1.8 * cm, 0.92 * cm, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf() -> None:
    doc = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.9 * cm,
        bottomMargin=1.7 * cm,
        title="Smart Logistics Management System Final Submission",
        author="Prachyat Misra, Sujal Purbey, Sudhanshu Ranjan",
    )
    story = build_story()
    doc.build(story, onFirstPage=draw_pdf_header_footer, onLaterPages=draw_pdf_header_footer)


def set_docx_page_border(section) -> None:
    sect_pr = section._sectPr
    pg_borders = sect_pr.xpath("./w:pgBorders")
    if pg_borders:
        pg_borders = pg_borders[0]
    else:
        pg_borders = sect_pr.makeelement(qn("w:pgBorders"))
        sect_pr.append(pg_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = pg_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = sect_pr.makeelement(qn(f"w:{edge}"))
            pg_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "12")
        element.set(qn("w:space"), "24")
        element.set(qn("w:color"), "214E7A")


def add_docx_heading(document: Document, text: str, level: int = 1, center: bool = False):
    paragraph = document.add_heading(text, level=level)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def add_docx_paragraph(document: Document, text: str, bold: bool = False, center: bool = False, font_size: int = 11):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(font_size)
    return paragraph


def add_docx_code(document: Document, code: str):
    for line in code.splitlines():
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(7.5)


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    set_docx_page_border(section)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(11)

    add_docx_paragraph(document, "PROGRESS REVIEW 3 / FINAL SUBMISSION", bold=True, center=True, font_size=14)
    add_docx_paragraph(document, "SMART LOGISTICS MANAGEMENT SYSTEM", bold=True, center=True, font_size=20)
    add_docx_paragraph(document, "(SMLOG-JAVA)", bold=True, center=True, font_size=14)
    add_docx_paragraph(
        document,
        "Servlet-based shipment tracking platform with MySQL persistence, browser dashboard, JavaFX desktop client, and Python machine learning integration.",
        center=True,
    )
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Report Type"
    hdr[1].text = "Final Project Review / DA2 Submission"
    for label, value in [
        ("Prepared For", "Faculty review submission based on the original Smart Logistics template"),
        ("Implementation Stage", "Project completed and consolidated for final submission"),
        ("Date", "15 April 2026"),
    ]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    document.add_page_break()
    add_docx_heading(document, "Project Overview", level=1, center=True)
    add_docx_paragraph(
        document,
        "Smart Logistics Management System is a multi-layer Java project designed to manage shipments, expose servlet-backed APIs, persist logistics data in MySQL, offer both browser and JavaFX interfaces, and provide AI-assisted forecasting and anomaly detection through Python integration.",
    )
    add_docx_heading(document, "Team Contributions", level=1, center=True)
    member_table = document.add_table(rows=1, cols=3)
    member_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    member_table.rows[0].cells[0].text = "Team Member"
    member_table.rows[0].cells[1].text = "Roll Number"
    member_table.rows[0].cells[2].text = "Primary Responsibility"
    for member in TEAM_MEMBERS:
        row = member_table.add_row().cells
        row[0].text = member["name"]
        row[1].text = member["roll"]
        row[2].text = member["responsibility"]

    add_docx_heading(document, "Architecture", level=1, center=True)
    architecture_png = ensure_architecture_png()
    document.add_picture(str(architecture_png), width=Inches(6.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_docx_heading(document, "Execution Evidence", level=1, center=True)
    for title, path, caption in SCREENSHOTS:
        add_docx_heading(document, title, level=2)
        document.add_picture(str(path), width=Inches(6.0))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_docx_paragraph(document, caption, center=True, font_size=9)

    document.add_page_break()
    add_docx_heading(document, "Full Source Code Listings", level=1, center=True)
    for rel_path in CORE_FILES:
        add_docx_heading(document, rel_path, level=2)
        add_docx_paragraph(
            document,
            f"Language: {language_for_path(rel_path)} | Lines: {count_loc(rel_path)} | Owner: {owner_for_path(rel_path)}",
            font_size=9,
        )
        add_docx_code(document, load_code(rel_path))
        document.add_page_break()

    document.save(str(OUTPUT_DOCX))


def main() -> None:
    build_pdf()
    build_docx()
    print(f"Created PDF: {OUTPUT_PDF}")
    print(f"Created DOCX: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
